from docx import Document

doc = Document()

doc.add_heading("AI Generated Document", 0)

doc.add_paragraph("This document is generated using Python.")
doc.add_paragraph("Next step: integrate OpenAI.")

doc.save("output.docx")

print("Document created!")

