import streamlit as st
from docx import Document
from pathlib import Path

st.title("AI Document Generator Prototype")

Path("output").mkdir(exist_ok=True)

user_text = st.text_area("Paste company notes here")
uploaded_file = st.file_uploader("Upload a file", type=["txt"])

if st.button("Generate"):
    final_text = user_text.strip()

    if uploaded_file is not None:
        file_content = uploaded_file.read().decode("utf-8")
        final_text += "\n\n--- Uploaded File Content ---\n" + file_content

    if not final_text.strip():
        st.warning("Please enter text or upload a file.")
    else:
        summary = f"""Structured Summary

Input received successfully.

Key content:
{final_text}

Status:
Document generated without OpenAI API because API quota is not available yet.
"""

        st.subheader("Output")
        st.text(summary)

        txt_path = "output/generated_summary.txt"
        docx_path = "output/generated_summary.docx"

        with open(txt_path, "w") as f:
            f.write(summary)

        doc = Document()
        doc.add_heading("Generated Business Document", 0)
        doc.add_paragraph(summary)
        doc.save(docx_path)

        st.success("TXT and DOCX files generated successfully.")
